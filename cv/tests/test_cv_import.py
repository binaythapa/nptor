from io import BytesIO

from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase
from django.urls import reverse

from cv.models import CareerEducation, CareerExperience, CareerProject, CareerSkill
from cv.models_import import ImportedField
from cv.services.importers.docx import extract_text_from_docx
from cv.services.importers.parser import parse_career_facts
from cv.services.importers.pdf import extract_text_from_pdf
from cv.services.importers.service import confirm_import_field, import_cv_source


class CVImportTests(TestCase):
    def test_common_fields_are_parsed_as_unconfirmed(self):
        result = parse_career_facts(
            "John Doe\nSoftware Engineer\njohn@example.com\n"
            "Skills: Python, Django, SQL"
        )
        self.assertEqual(result["full_name"], "John Doe")
        self.assertEqual(result["email"], "john@example.com")
        self.assertFalse(result["fields"][0]["confirmed"])

    def test_section_heading_is_not_used_as_name_or_professional_title(self):
        result = parse_career_facts(
            "Summary\n"
            "Gyanendra Thapa is a Software Engineer, specializing in Business Intelligence (BI).\n"
            "Experience\n"
            "Accenture | Business Intelligence"
        )

        self.assertEqual(result["full_name"], "Gyanendra Thapa")
        self.assertEqual(result["professional_title"], "Software Engineer")
        self.assertNotEqual(result["full_name"], "Summary")
        self.assertNotEqual(result["professional_title"], "Gyanendra Thapa is a Software Engineer, specializing in Business Intelligence (BI).")
        summary_field = next(field for field in result["fields"] if field["section"] == "summary")
        self.assertEqual(summary_field["field_name"], "text")

    def test_inline_resume_sections_are_split_instead_of_becoming_one_summary(self):
        result = parse_career_facts(
            "Gyanendra Thapa | Software Engineer | gyanendra@example.com "
            "Professional Summary: Software engineer with BI and ETL experience. "
            "Work Experience: Software Engineer | Accenture | Built ETL pipelines. "
            "Education: B.Tech | Computer Science | ABC University. "
            "Technical Skills: Python, SQL, Snowflake, Power BI. "
            "Projects: Retail Analytics Platform | Built a reporting platform."
        )

        sections = [field["section"] for field in result["fields"]]
        self.assertIn("summary", sections)
        self.assertIn("experience", sections)
        self.assertIn("education", sections)
        self.assertIn("skills", sections)
        self.assertIn("projects", sections)
        summary = next(field["value"] for field in result["fields"] if field["section"] == "summary")
        self.assertNotIn("Accenture", summary)
        self.assertNotIn("ABC University", summary)
        self.assertNotIn("Snowflake", summary)

    def test_docx_heading_styles_are_preserved_for_import_parsing(self):
        from docx import Document

        stream = BytesIO()
        document = Document()
        document.add_paragraph("John Doe")
        document.add_paragraph("Software Engineer")
        document.add_paragraph("Professional Summary", style="Heading 1")
        document.add_paragraph("Engineer with ETL experience.")
        document.add_paragraph("Work Experience", style="Heading 1")
        document.add_paragraph("Software Engineer | Accenture")
        document.add_paragraph("Education", style="Heading 1")
        document.add_paragraph("B.Tech | ABC University")
        document.add_paragraph("Technical Skills", style="Heading 1")
        document.add_paragraph("Python, SQL, Snowflake")
        document.save(stream)
        stream.seek(0)

        uploaded = SimpleUploadedFile(
            "resume.docx",
            stream.read(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        text = extract_text_from_docx(uploaded)
        result = parse_career_facts(text)

        self.assertIn("Professional Summary", text)
        self.assertIn("Work Experience", text)
        self.assertIn("Education", text)
        self.assertIn("Technical Skills", text)
        self.assertTrue(any(field["section"] == "experience" for field in result["fields"]))
        self.assertTrue(any(field["section"] == "education" for field in result["fields"]))
        self.assertTrue(any(field["section"] == "skills" for field in result["fields"]))

    def test_import_review_uses_section_cards_and_textarea_for_text_fields(self):
        imported = import_cv_source(self.user, self._pdf_upload("John Doe", "john@example.com"))
        summary = ImportedField.objects.create(
            cv_import=imported,
            section="summary",
            field_name="text",
            value="John Doe is a Software Engineer with experience in ETL and Business Intelligence.",
        )

        self.client.force_login(self.user)
        response = self.client.get(reverse("cv:cv_import_review", args=[imported.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Review your imported CV")
        self.assertContains(response, "Contact")
        self.assertContains(response, "Summary")
        self.assertContains(response, "Full Name")
        self.assertContains(response, "<textarea", html=False)
        self.assertContains(response, f'name="field_{summary.pk}"', html=False)
        self.assertContains(response, "Please verify")
        self.assertContains(response, "Confirm &amp; Add to CV", html=False)

    def test_confirm_and_add_to_cv_applies_imported_profile_fields(self):
        imported = import_cv_source(self.user, self._pdf_upload("John Doe", "john@example.com"))
        title = ImportedField.objects.create(
            cv_import=imported,
            section="contact",
            field_name="professional_title",
            value="Business Intelligence Engineer",
        )
        summary = ImportedField.objects.create(
            cv_import=imported,
            section="summary",
            field_name="text",
            value="Experienced BI engineer specializing in ETL and analytics.",
        )
        skill = ImportedField.objects.create(
            cv_import=imported,
            section="skills",
            field_name="name",
            value="Snowflake",
        )

        self.client.force_login(self.user)
        response = self.client.post(
            reverse("cv:cv_import_review", args=[imported.pk]),
            {
                f"field_{field.pk}": field.value
                for field in imported.fields.all()
            },
        )

        self.assertRedirects(response, reverse("cv:dashboard"))
        self.user.refresh_from_db()
        imported.refresh_from_db()
        profile = imported.profile
        profile.refresh_from_db()
        self.assertEqual(self.user.first_name, "John")
        self.assertEqual(self.user.last_name, "Doe")
        self.assertEqual(self.user.email, "john@example.com")
        self.assertEqual(profile.professional_title, title.value)
        self.assertEqual(profile.summary, summary.value)
        self.assertTrue(profile.careerskill_records.filter(name=skill.value).exists())
        self.assertEqual(imported.status, imported.STATUS_CONFIRMED)
        self.assertFalse(imported.fields.filter(confirmed=False).exists())

    def test_confirm_import_creates_experience_education_and_project_records(self):
        imported = import_cv_source(self.user, self._pdf_upload("John Doe", "john@example.com"))
        fields = [
            ImportedField.objects.create(
                cv_import=imported,
                section="experience",
                field_name="text",
                value="Software Engineer | Accenture\nBuilt ETL pipelines and BI reports.",
            ),
            ImportedField.objects.create(
                cv_import=imported,
                section="education",
                field_name="text",
                value="B.Tech | Computer Science | ABC University",
            ),
            ImportedField.objects.create(
                cv_import=imported,
                section="projects",
                field_name="text",
                value="Retail Analytics Platform\nBuilt a reporting platform using Snowflake and Power BI.",
            ),
        ]

        self.client.force_login(self.user)
        response = self.client.post(
            reverse("cv:cv_import_review", args=[imported.pk]),
            {f"field_{field.pk}": field.value for field in imported.fields.all()},
        )

        self.assertRedirects(response, reverse("cv:dashboard"))
        profile = imported.profile
        self.assertTrue(profile.careerexperience_records.filter(job_title="Software Engineer", employer="Accenture").exists())
        self.assertTrue(profile.careereducation_records.filter(qualification="B.Tech", institution="ABC University").exists())
        self.assertTrue(profile.careerproject_records.filter(name="Retail Analytics Platform").exists())
        self.assertEqual(len(fields), 3)

    def test_pdf_adapter_extracts_text(self):
        from reportlab.pdfgen import canvas

        stream = BytesIO()
        pdf = canvas.Canvas(stream)
        pdf.drawString(72, 720, "John Doe")
        pdf.save()
        stream.seek(0)

        uploaded = SimpleUploadedFile("resume.pdf", stream.read(), content_type="application/pdf")
        self.assertIn("John Doe", extract_text_from_pdf(uploaded))

    def test_docx_adapter_extracts_text(self):
        from docx import Document

        stream = BytesIO()
        document = Document()
        document.add_paragraph("John Doe")
        document.add_paragraph("Software Engineer")
        document.save(stream)
        stream.seek(0)

        uploaded = SimpleUploadedFile(
            "resume.docx",
            stream.read(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        self.assertIn("John Doe", extract_text_from_docx(uploaded))

    def test_unsupported_format_is_rejected(self):
        uploaded = SimpleUploadedFile("resume.txt", b"John Doe", content_type="text/plain")
        with self.assertRaises(ValueError):
            import_cv_source(self.user, uploaded)

    def test_import_stores_unconfirmed_fields_and_owner(self):
        imported = import_cv_source(self.user, self._pdf_upload("John Doe", "john@example.com"))

        self.assertEqual(imported.owner_id, self.user.id)
        self.assertEqual(imported.status, imported.STATUS_REVIEW)
        self.assertTrue(imported.extracted_text)
        self.assertTrue(imported.fields.filter(confirmed=False).exists())

    def test_confirm_import_field_requires_owner_and_marks_confirmed(self):
        imported = import_cv_source(self.user, self._pdf_upload("John Doe", "john@example.com"))
        field = imported.fields.get(field_name="full_name")
        confirmed = confirm_import_field(field.pk, self.user, "John A. Doe")

        self.assertTrue(confirmed.confirmed)
        self.assertEqual(confirmed.confirmed_by_id, self.user.id)
        self.assertEqual(confirmed.value, "John A. Doe")

    def test_confirm_import_field_cannot_be_used_by_another_user(self):
        from django.contrib.auth import get_user_model

        imported = import_cv_source(self.user, self._pdf_upload("John Doe", "john@example.com"))
        field = imported.fields.get(field_name="full_name")
        other_user = get_user_model().objects.create_user(username="other-cv-user", password="password")

        with self.assertRaises(ImportedField.DoesNotExist):
            confirm_import_field(field.pk, other_user, "Changed")

        field.refresh_from_db()
        self.assertFalse(field.confirmed)

    def _pdf_upload(self, name, email):
        from reportlab.pdfgen import canvas

        stream = BytesIO()
        pdf = canvas.Canvas(stream)
        pdf.drawString(72, 720, name)
        pdf.drawString(72, 700, email)
        pdf.save()
        stream.seek(0)
        return SimpleUploadedFile("resume.pdf", stream.read(), content_type="application/pdf")

    def setUp(self):
        from django.contrib.auth import get_user_model

        self.user = get_user_model().objects.create_user(
            username="cv-import-user", email="cv@example.com", password="password"
        )
