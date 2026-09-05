from io import BytesIO

from django.core.files.base import ContentFile
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from cv.models_document import DocumentArtifact
from cv.services.documents.base import DocumentGenerator
from cv.services.documents.renderer import get_render_config


MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
FONT_MAP = {
    "Helvetica": "Arial",
    "Helvetica-Bold": "Arial",
    "Times-Roman": "Times New Roman",
    "Times-Bold": "Times New Roman",
    "Courier": "Courier New",
}


def _rgb(hex_color):
    value = hex_color.lstrip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def _set_cell_shading(cell, hex_color):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color.lstrip("#"))
    cell._tc.get_or_add_tcPr().append(shading)


def _remove_table_borders(table):
    borders = table._tbl.tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table._tbl.tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "nil")


class DOCXGenerator(DocumentGenerator):
    def generate(self, version):
        payload = version.snapshot
        config = get_render_config(payload.get("template", {}))
        document = Document()
        section = document.sections[0]
        margin = Inches(config["margin"] / 72)
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin

        normal = document.styles["Normal"]
        font_name = FONT_MAP.get(config["font_name"], "Arial")
        normal.font.name = font_name
        normal.font.size = Pt(config["font_size"])
        accent = _rgb(config["accent_color"])
        accent_hex = config["accent_color"].lstrip("#")
        compact = config["density"] == "compact"

        contact = payload.get("contact", {})
        name = f"{contact.get('first_name', '')} {contact.get('last_name', '')}".strip()
        details = [contact.get("email"), contact.get("phone"), contact.get("location")]
        details = [str(value) for value in details if value]
        if payload.get("linkedin_url"):
            details.append(payload["linkedin_url"])
        if payload.get("portfolio_url"):
            details.append(payload["portfolio_url"])

        header = document.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER if config["header_style"] == "centered" else WD_ALIGN_PARAGRAPH.LEFT
        if name:
            run = header.add_run(name)
            run.bold = True
            run.font.name = font_name
            run.font.size = Pt(config["heading_size"] + (5 if config["header_style"] == "compact" else 8))
            run.font.color.rgb = accent
        if payload.get("professional_title"):
            run = header.add_run(f"\n{payload['professional_title']}")
            run.font.name = font_name
            run.font.size = Pt(config["font_size"] + 2)
            run.font.color.rgb = accent
        if details:
            run = header.add_run(f"\n{' | '.join(details)}")
            run.font.name = font_name
            run.font.size = Pt(max(8, config["font_size"] - 1))

        if config["header_style"] != "compact":
            rule = document.add_paragraph()
            rule.paragraph_format.space_after = Pt(4)
            run = rule.add_run("_" * 95)
            run.font.color.rgb = accent

        def add_heading(title, target, white=False):
            paragraph = target.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(4 if compact else 7)
            paragraph.paragraph_format.space_after = Pt(2 if compact else 4)
            text = title.upper() if config["section_style"] == "uppercase_rule" else title
            run = paragraph.add_run(text)
            run.bold = True
            run.font.name = font_name
            run.font.size = Pt(config["heading_size"])
            run.font.color.rgb = RGBColor(255, 255, 255) if white else accent
            if config["section_style"] == "uppercase_rule" and not white:
                paragraph.add_run("\n" + "─" * 55)
            return paragraph

        def add_item(target, kind, item, white=False):
            if kind == "experience":
                title = f"{item.get('job_title', '')} — {item.get('employer', '')}"
                body = item.get("description", "")
            elif kind == "education":
                title = f"{item.get('qualification', '')} — {item.get('institution', '')}"
                body = item.get("field_of_study", "")
            elif kind == "skills":
                title = item.get("name", "")
                body = item.get("proficiency", "")
            elif kind == "projects":
                title = f"{item.get('name', '')} — {item.get('role', '')}"
                body = item.get("description", "")
            elif kind == "certifications":
                title = f"{item.get('name', '')} — {item.get('issuer', '')}"
                body = item.get("credential_id", "")
            else:
                title = item.get("title", "")
                body = item.get("description", "")
            paragraph = target.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(2 if compact else 4)
            if title:
                run = paragraph.add_run(str(title))
                run.bold = True
                run.font.name = font_name
                run.font.size = Pt(max(8, config["font_size"] - 1))
                if white:
                    run.font.color.rgb = RGBColor(255, 255, 255)
            if body:
                run = paragraph.add_run(f"\n{body}")
                run.font.name = font_name
                run.font.size = Pt(max(8, config["font_size"] - 1))
                if white:
                    run.font.color.rgb = RGBColor(245, 245, 245)

        def render_group(target, heading, value, kind, white=False):
            if not value:
                return
            add_heading(heading, target, white=white)
            if kind == "text":
                paragraph = target.add_paragraph(str(value))
                if white:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(245, 245, 245)
                return
            if kind == "skills":
                for item in value:
                    add_item(target, kind, item, white=white)
                return
            for item in value:
                add_item(target, kind, item, white=white)

        sidebar_groups = (
            ("Skills", payload.get("skills", []), "skills"),
            ("Education", payload.get("educations", []), "education"),
            ("Certifications", payload.get("certifications", []), "certifications"),
        )
        main_groups = (
            ("Summary", payload.get("summary"), "text"),
            ("Experience", payload.get("experiences", []), "experience"),
            ("Projects", payload.get("projects", []), "projects"),
            ("Achievements", payload.get("achievements", []), "achievements"),
        )

        if config["layout"] == "sidebar":
            table = document.add_table(rows=1, cols=2)
            table.autofit = False
            _remove_table_borders(table)
            sidebar_cell, main_cell = table.rows[0].cells
            sidebar_cell.width = Inches(2.0)
            main_cell.width = Inches(4.8)
            sidebar_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            main_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            _set_cell_shading(sidebar_cell, accent_hex)
            for heading, value, kind in sidebar_groups:
                render_group(sidebar_cell, heading, value, kind, white=True)
            for heading, value, kind in main_groups:
                render_group(main_cell, heading, value, kind)
        else:
            for heading, value, kind in main_groups:
                render_group(document, heading, value, kind)
            for heading, value, kind in sidebar_groups:
                render_group(document, heading, value, kind)

        stream = BytesIO()
        document.save(stream)
        stream.seek(0)
        artifact = DocumentArtifact(
            cv_version=version,
            artifact_type=DocumentArtifact.DOCX,
            mime_type=MIME_TYPE,
            template_slug=payload["template"]["slug"],
            template_config=payload["template"].get("config", {}),
        )
        artifact.file.save(f"{version.cv_id}-v{version.version_number}.docx", ContentFile(stream.read()), save=True)
        return artifact


def generate_docx(cv_version):
    return DOCXGenerator().generate(cv_version)
