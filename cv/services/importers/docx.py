from docx import Document

from cv.services.importers.base import CVImportAdapter


def _is_heading_paragraph(paragraph):
    style_name = (getattr(getattr(paragraph, "style", None), "name", "") or "").strip().lower()
    return style_name == "title" or style_name.startswith("heading") or "heading" in style_name


def extract_text_from_docx(uploaded_file):
    uploaded_file.seek(0)
    document = Document(uploaded_file)
    parts = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        # Keep Word heading paragraphs visually separated so the parser can reliably
        # identify section boundaries even when the original document has complex styling.
        if _is_heading_paragraph(paragraph):
            parts.append(f"\n{text}\n")
        else:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    uploaded_file.seek(0)
    return "\n".join(parts).strip()


class DOCXImportAdapter(CVImportAdapter):
    source_type = "docx"
    extract_text = staticmethod(extract_text_from_docx)
